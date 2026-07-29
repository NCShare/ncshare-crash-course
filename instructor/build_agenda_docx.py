#!/usr/bin/env python3
"""Build the concise NCShare agenda as a styled Word document."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "agenda" / "Agenda_Revised.docx"
TABLE_HELPER = Path(
    "/Users/alejo/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.727.11326/skills/documents/scripts"
)
sys.path.insert(0, str(TABLE_HELPER))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "2F3337"
MUTED = "5D6772"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "B8C4D1"


def set_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)
        borders.append(elem)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.5, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    run.append(props)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(text)
    run.append(end)
    paragraph._p.append(run)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(tabs)
    p_pr.append(ind)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    for node in (start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr):
        level.append(node)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_bullet(doc, text, num_id):
    paragraph = doc.add_paragraph()
    apply_bullet(paragraph, num_id)
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_body(doc, text, *, bold_lead=None, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_font(lead, size=11, bold=True, color=NAVY)
        body = paragraph.add_run(text[len(bold_lead) :])
        set_font(body, size=11)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11)
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=10.5, color=NAVY, bold=True)
    text_run = paragraph.add_run(text)
    set_font(text_run, size=10.5, color=INK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    values = [
        ("DATE", "August 19"),
        ("FORMAT", "In person"),
        ("HOURS", "9:00-5:00"),
        ("AUDIENCE", "New HPC users"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(label + "\n")
        set_font(label_run, size=8.5, color=BLUE, bold=True)
        value_run = paragraph.add_run(value)
        set_font(value_run, size=10.5, color=NAVY, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [2340, 2340, 2340, 2340])
    set_table_borders(table, color="D4DDE7", size="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_schedule_table(doc):
    rows = [
        ("9:00-9:30", "Check-in and setup", "Verified access and course clone"),
        ("9:30-9:45", "Welcome and goals", "Shared runnable outcomes"),
        ("9:45-10:30", "Session 1A: cluster model", "Login/storage/module preflight"),
        ("10:30-10:45", "Break", ""),
        ("10:45-11:30", "Session 1B: guided practice", "First interactive allocation"),
        ("11:30-12:00", "Session 2: storage and I/O", "Data-lifecycle plan"),
        ("12:00-1:00", "Lunch and discussion", ""),
        ("1:00-2:15", "Session 3A: inoisy+ on CPUs", "One-rank/four-rank HDF5 results"),
        ("2:15-2:30", "Break", ""),
        ("2:30-3:30", "Session 3B: QuantUI on GPU", "Verified GPU-offload result"),
        ("3:30-4:30", "Session 4: visualization", "Post-processed figure and notebook"),
        ("4:30-5:00", "Wrap-up", "Support path and next steps"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["Time", "Session", "Participant output"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=NAVY, size=9.5)
    set_repeat_table_header(table.rows[0])
    for time, session, output in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], time, bold=True, color=DARK_BLUE, size=9.2)
        set_cell_text(cells[1], session, size=9.2)
        set_cell_text(cells[2], output, size=9.2)
        if session in {"Break", "Lunch and discussion"}:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
    apply_table_geometry(table, [1680, 3180, 4500])
    set_table_borders(table)


def add_session_plan_table(doc):
    rows = [
        ("9:45-9:55", "Map laptop → login → Slurm → compute → storage.", "Annotated workflow"),
        ("9:55-10:10", "Explain CPU/GPU resources, partitions, memory, and wall time.", "Resource vocabulary"),
        ("10:10-10:20", "Log in; verify identity, host, quota, and storage.", "Completed preflight"),
        ("10:20-10:30", "Use ten commands and inspect the module system.", "Command/module record"),
        ("10:30-10:45", "Break.", ""),
        ("10:45-11:00", "Clone the course and NCShare examples.", "Local repositories"),
        ("11:00-11:20", "Request a compute shell; compare hosts and allocation variables.", "First allocation"),
        ("11:20-11:30", "Answer: where am I, what do I have, where does output go?", "Checkpoint answers"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for cell, header in zip(table.rows[0].cells, ("Time", "Teach / do", "Concrete output")):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=NAVY, size=9.5)
    set_repeat_table_header(table.rows[0])
    for time, activity, output in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], time, bold=True, color=DARK_BLUE, size=9.2)
        set_cell_text(cells[1], activity, size=9.2)
        set_cell_text(cells[2], output, size=9.2)
        if activity == "Break.":
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
    apply_table_geometry(table, [1500, 4650, 3210])
    set_table_borders(table)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("NCShare Crash Course")
    set_font(left, size=9, color=MUTED, bold=True)
    right = p.add_run("\tWorkshop Agenda")
    set_font(right, size=9, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    label = p.add_run("Page ")
    set_font(label, size=9, color=MUTED)
    add_page_field(p)


def build():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_num = configure_numbering(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(0)
    run = kicker.add_run("WORKSHOP AGENDA")
    set_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("NCShare Crash Course")
    set_font(run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "From cluster access to a reproducible CPU, GPU, and visualization workflow"
    )
    set_font(run, size=13.2, color=MUTED)

    add_metric_strip(doc)

    add_heading(doc, "Purpose", 1)
    add_body(
        doc,
        "Move participants from “I have access” to “I can choose resources, submit a job, "
        "inspect its output, and continue my analysis.” The day uses the official NCShare "
        "guides and examples throughout.",
    )
    links = doc.add_paragraph()
    links.paragraph_format.space_after = Pt(6)
    add_hyperlink(links, "NCShare user guides", "https://userguide.ncshare.org/guides/")
    links.add_run("  •  ")
    add_hyperlink(links, "NCShare examples", "https://github.com/NCShare/examples")

    add_callout(
        doc,
        "Before the workshop",
        "Active NCShare account; registered SSH public key; laptop with SSH client; "
        "course clone; user-owned Miniforge; GPU access requested in advance. Pairing "
        "is available for participants still waiting for access.",
    )

    add_heading(doc, "Learning outcomes", 1)
    outcomes = [
        "Distinguish login, compute, CPU, and GPU nodes.",
        "Choose among /hpc/home, /work, and job-local /scratch.",
        "Load modules, install a user-space library, compile C/MPI code, and manage conda.",
        "Submit, monitor, diagnose, and cancel Slurm jobs.",
        "Run one-rank, four-rank MPI, and GPU workflows with reasonable resources.",
        "Inspect HDF5, post-process a GRF, and export an accessible scientific figure.",
    ]
    for item in outcomes:
        add_bullet(doc, item, bullet_num)

    doc.add_page_break()
    add_heading(doc, "At a glance", 1)
    add_schedule_table(doc)

    add_heading(doc, "Session 1 • Access, cluster model, and essential tools", 1)
    add_body(
        doc,
        "9:45-11:30 (with a 10:30-10:45 break). Explain only what participants need "
        "to request their first allocation, then practice immediately.",
    )
    add_session_plan_table(doc)
    add_callout(
        doc,
        "HPC administrator contribution",
        "HPC admins provide workshop access and teach the NCShare login/compute boundary, "
        "storage lifetimes, partitions, and support path.",
    )
    add_body(
        doc,
        "Provided: command card, annotated cluster workflow, access troubleshooting, "
        "and a verified interactive-allocation command.",
    )

    add_heading(doc, "Session 2 • Storage, transfer, and I/O", 1)
    add_body(
        doc,
        "11:30-12:00. Participants choose the correct location for code, active data, "
        "temporary I/O, and retained results; practice scp/rsync; and write the data "
        "lifecycle for the afternoon jobs.",
    )
    for item in (
        "Use /hpc/home for scripts, environments, and user software.",
        "Use /work for active data and move retained results off NCShare.",
        "Use job-local /scratch for temporary high-I/O work and copy results out.",
        "Never place sensitive data on NCShare.",
    ):
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "Session 3 • From source code to CPU and GPU jobs", 1)
    add_body(
        doc,
        "1:00-3:30 (with a 2:15-2:30 break). This merged block covers clone → "
        "environment → install/compile → request → submit → monitor → inspect.",
    )

    add_heading(doc, "1:00-1:45 • inoisy+ on CPUs", 2)
    for item in (
        "Clone a real scientific C/MPI repository and inspect its README and Makefile.",
        "Load compiler, MPI, parallel HDF5, and GSL modules.",
        "Install a user-space HYPRE build with four-dimensional SStruct support.",
        "Compile unmodified inoisy4d without editing its source or Makefile.",
        "Submit the same 16⁴ global grid with one and four MPI ranks; retain HDF5 output.",
    ):
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "1:45-2:15 • Scheduler and efficiency debrief", 2)
    add_body(
        doc,
        "Use squeue, logs, and sacct to compare job states, elapsed time, memory, and "
        "exit codes. Treat slower parallel timing for a tiny problem as a lesson in "
        "overhead, not a reason to request more resources.",
    )

    doc.add_page_break()
    add_heading(doc, "2:30-3:15 • QuantUI on an H200 GPU", 2)
    for item in (
        "Create a Python 3.11 conda environment on a CPU allocation.",
        "Install QuantUI plus CUDA 12.x GPU wheels and register its Jupyter kernel.",
        "Request one H200 GPU, verify the device, and submit a small RHF calculation.",
        "Confirm QuantUI reports gpu_used=true; do not infer use from allocation alone.",
    ):
        add_bullet(doc, item, bullet_num)

    add_heading(doc, "3:15-3:30 • CPU/GPU comparison", 2)
    add_body(
        doc,
        "Compare the Slurm files, identify work that does not benefit from a GPU, and "
        "record one resource change to make before scaling.",
    )
    add_callout(
        doc,
        "Provided",
        "HYPRE/build helpers, one-rank/four-rank/GPU Slurm files, low-resolution "
        "settings, expected outputs, and troubleshooting checkpoints.",
    )

    add_heading(doc, "Session 4 • Scientific visualization and post-processing", 1)
    add_body(
        doc,
        "3:30-4:30. In Jupyter, inspect HDF5 lazily, plot distributions and slices, "
        "choose honest color/normalization, run the upstream GRF-to-emissivity converter, "
        "compare raw and positive fields, and export PNG/PDF figures with provenance.",
    )

    add_heading(doc, "Wrap-up • 4:30-5:00", 1)
    add_body(
        doc,
        "Review the end-to-end workflow, locate NCShare documentation and local support, "
        "capture unresolved questions, and show how participants can contribute examples.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
